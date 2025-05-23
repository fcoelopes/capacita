import os
import time
import pandas as pd
from docx import Document
from watchdog.observers import Observer
from watchdog.events import FileSystemEventHandler

# Caminhos
raw_folder = './data/raw'
preprocessed_folder = './data/preprocessed'
processed_folder = './data/processed'

paths = [raw_folder, preprocessed_folder, processed_folder]
for path in paths:
    if not os.path.exists(path):
        os.makedirs(path)

class WordFileHandler(FileSystemEventHandler):
    def on_created(self, event):
        if not event.is_directory and event.src_path.endswith('.docx'):
            time.sleep(1)  # Pequeno atraso para garantir que o arquivo foi completamente escrito
            processar_documento(event.src_path)

def extrair_conteudo(caminho):
    doc = Document(caminho)
    tables = []
    texts = []

    for element in doc.element.body:
        if element.tag.endswith('}p'):  # Parágrafo
            para = element
            paragraph = [p for p in doc.paragraphs if p._p == para]
            if paragraph:
                texts.append(paragraph[0].text)
        elif element.tag.endswith('}tbl'):  # Tabela
            tbl = element
            table = [t for t in doc.tables if t._tbl == tbl]
            if table:
                # Extrair os dados da tabela como lista de listas
                temp_table = []
                for row in table[0].rows:
                    linha = [cell.text.strip() for cell in row.cells]
                    temp_table.append(linha)
                df = pd.DataFrame(temp_table)
                if not df.empty:
                    df.columns = df.iloc[0] # Definindo a primeira linha como cabeçalho
                    df = df[1:].reset_index(drop=True) # Removendo a primeira linha que agora é o cabeçalho
                tables.append(df)
    return texts, tables


def salvar_tabelas_em_csv(tabelas, nome_arquivo_base, metodo):
    for i, df in enumerate(tabelas):
        filename = f"{nome_arquivo_base}_tabela_{i+1}_{metodo}.csv"
        preprocessed_path = os.path.join(preprocessed_folder, filename)
        df.to_csv(preprocessed_path, index=False)

def salvar_textos_em_txt(textos, nome_arquivo_base):
    pass

def processar_documento(caminho):
    arquivo = os.path.basename(caminho)
    nome_base = os.path.splitext(arquivo)[0]
    _, tabelas_docx = extrair_conteudo(caminho) # Extrai textos e tabelas
    salvar_tabelas_em_csv(tabelas_docx, nome_base, 'python_docx')

class CSVFileHandler(FileSystemEventHandler):
    def on_created(self, event):
        if not event.is_directory and event.src_path.endswith('.csv'):
            time.sleep(1)  # Pequeno atraso para garantir que o arquivo foi completamente escrito
            unir_arquivos_csv(preprocessed_folder)


def unir_arquivos_csv(pasta):
    arquivos_csv = [os.path.join(pasta, f) for f in os.listdir(pasta) if f.endswith('.csv')]
    dataframes = []
    
    for arquivo in arquivos_csv:
        try:
            df = pd.read_csv(arquivo)
            dataframes.append(df)
        except Exception as e:
            print(f"Erro ao ler o arquivo {arquivo}: {e}")
    
    if dataframes:
        df_unido = pd.concat(dataframes, ignore_index=True)
        caminho_saida = os.path.join(processed_folder, 'tabelas.csv')
        df_unido.to_csv(caminho_saida, index=False)
        print(f"Arquivo salvo em {caminho_saida}")
    else:
        print("Nenhum arquivo CSV encontrado para unir.")

def iniciar_monitoramento():
    observer_raw = Observer()
    observer_raw.schedule(WordFileHandler(), path=raw_folder, recursive=True)
    observer_raw.start()

    observer_preprocessed = Observer()
    observer_preprocessed.schedule(CSVFileHandler(), path=preprocessed_folder, recursive=True)
    observer_preprocessed.start()

    print(f"Monitorando pastas: {raw_folder}, {preprocessed_folder}")
    print("Pressione Ctrl+C para parar o monitoramento.")

    try:
        while True:
            time.sleep(1)
    except KeyboardInterrupt:
        observer_raw.stop()
        observer_preprocessed.stop()
    observer_raw.join()
    observer_preprocessed.join()


if __name__ == '__main__':
    iniciar_monitoramento()
